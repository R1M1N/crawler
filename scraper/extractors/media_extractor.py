"""
Media Content Extractor

Handles extraction and processing of media content including:
- Images (download, analyze, OCR)
- Videos (download, extract frames, transcribe audio)
- Audio files (download, transcribe)
- Documents (PDFs, DOCX with images)
"""

import asyncio
import aiohttp
import hashlib
import logging
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Any, Tuple
from urllib.parse import urljoin, urlparse
import json

from ..utils.logger import get_logger
from ..processors.ocr_processor import OCRProcessor
from ..processors.transcription_processor import TranscriptionProcessor

logger = get_logger(__name__)


class MediaExtractor:
    """Advanced media content extraction and processing."""
    
    def __init__(self, 
                 download_dir: Optional[Path] = None,
                 ocr_processor: Optional[OCRProcessor] = None,
                 transcription_processor: Optional[TranscriptionProcessor] = None):
        self.download_dir = download_dir or Path("media_downloads")
        self.download_dir.mkdir(exist_ok=True)
        
        self.ocr_processor = ocr_processor or OCRProcessor()
        self.transcription_processor = transcription_processor or TranscriptionProcessor()
        
        # Supported file types
        self.image_types = {'.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg'}
        self.video_types = {'.mp4', '.avi', '.mov', '.wmv', '.flv', '.webm', '.mkv', '.m4v'}
        self.audio_types = {'.mp3', '.wav', '.flac', '.aac', '.ogg', '.wma', '.m4a'}
        self.document_types = {'.pdf', '.docx', '.doc', '.txt', '.rtf'}
        
        # Session for downloads
        self.session: Optional[aiohttp.ClientSession] = None
        
    async def __aenter__(self):
        """Async context manager entry."""
        self.session = aiohttp.ClientSession()
        return self
        
    async def __aexit__(self, exc_type, exc_val, exc_tb):
        """Async context manager exit."""
        if self.session:
            await self.session.close()
            
    async def extract_media_from_page(self, 
                                    media_urls: List[str], 
                                    base_url: str = "",
                                    process_images: bool = True,
                                    process_videos: bool = True,
                                    process_audio: bool = True) -> Dict[str, Any]:
        """
        Extract and process all media from a page.
        
        Args:
            media_urls: List of media URLs to process
            base_url: Base URL for resolving relative URLs
            process_images: Whether to process images with OCR
            process_videos: Whether to process videos
            process_audio: Whether to process audio files
            
        Returns:
            Dictionary with processed media information
        """
        if not self.session:
            raise RuntimeError("MediaExtractor must be used as async context manager")
            
        media_results = {
            'images': [],
            'videos': [],
            'audio': [],
            'documents': [],
            'processed_at': datetime.now().isoformat()
        }
        
        # Categorize URLs by type
        categorized_urls = self._categorize_media_urls(media_urls, base_url)
        
        # Process each category
        if process_images and categorized_urls['images']:
            media_results['images'] = await self._process_images(categorized_urls['images'])
            
        if process_videos and categorized_urls['videos']:
            media_results['videos'] = await self._process_videos(categorized_urls['videos'])
            
        if process_audio and categorized_urls['audio']:
            media_results['audio'] = await self._process_audio(categorized_urls['audio'])
            
        if categorized_urls['documents']:
            media_results['documents'] = await self._process_documents(categorized_urls['documents'])
            
        return media_results
        
    def _categorize_media_urls(self, media_urls: List[str], base_url: str) -> Dict[str, List[str]]:
        """Categorize media URLs by file type."""
        categorized = {
            'images': [],
            'videos': [],
            'audio': [],
            'documents': []
        }
        
        for url in media_urls:
            # Resolve relative URLs
            if base_url and url.startswith('/'):
                url = urljoin(base_url, url)
                
            # Get file extension
            parsed_url = urlparse(url)
            path = parsed_url.path.lower()
            extension = Path(path).suffix
            
            # Categorize by extension
            if extension in self.image_types:
                categorized['images'].append(url)
            elif extension in self.video_types:
                categorized['videos'].append(url)
            elif extension in self.audio_types:
                categorized['audio'].append(url)
            elif extension in self.document_types:
                categorized['documents'].append(url)
            else:
                # Try to determine type from content-type if no extension
                categorized['documents'].append(url)  # Default to documents
                
        return categorized
        
    async def _process_images(self, image_urls: List[str]) -> List[Dict[str, Any]]:
        """Process images: download, analyze, and extract text."""
        processed_images = []
        
        # Limit concurrent downloads
        semaphore = asyncio.Semaphore(5)
        
        async def process_single_image(url: str):
            async with semaphore:
                try:
                    # Download image
                    local_path, image_info = await self._download_media(url)
                    
                    if not local_path:
                        return None
                        
                    # Analyze image
                    analysis = await self._analyze_image(local_path)
                    
                    # Extract text with OCR if possible
                    ocr_text = ""
                    if analysis.get('has_text', False):
                        try:
                            ocr_result = await self.ocr_processor.extract_text(str(local_path))
                            ocr_text = ocr_result.get('text', '')
                        except Exception as e:
                            logger.warning(f"OCR failed for {url}: {e}")
                            
                    return {
                        'url': url,
                        'local_path': str(local_path),
                        'file_size': image_info.get('size', 0),
                        'content_type': image_info.get('content_type', ''),
                        'dimensions': analysis.get('dimensions', {}),
                        'format': analysis.get('format', ''),
                        'has_text': analysis.get('has_text', False),
                        'dominant_colors': analysis.get('dominant_colors', []),
                        'ocr_text': ocr_text,
                        'processed_at': datetime.now().isoformat()
                    }
                    
                except Exception as e:
                    logger.error(f"Failed to process image {url}: {e}")
                    return None
                    
        # Process all images
        tasks = [process_single_image(url) for url in image_urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out failed results
        for result in results:
            if result and not isinstance(result, Exception):
                processed_images.append(result)
                
        return processed_images
        
    async def _process_videos(self, video_urls: List[str]) -> List[Dict[str, Any]]:
        """Process videos: download, extract metadata and audio."""
        processed_videos = []
        
        semaphore = asyncio.Semaphore(3)  # Limit concurrent video processing
        
        async def process_single_video(url: str):
            async with semaphore:
                try:
                    # Download video
                    local_path, video_info = await self._download_media(url)
                    
                    if not local_path:
                        return None
                        
                    # Extract video metadata
                    metadata = await self._extract_video_metadata(str(local_path))
                    
                    # Extract audio track for transcription
                    audio_transcription = ""
                    try:
                        audio_path = await self._extract_audio_from_video(str(local_path))
                        if audio_path:
                            transcription_result = await self.transcription_processor.transcribe_audio(str(audio_path))
                            audio_transcription = transcription_result.get('text', '')
                            # Clean up audio file
                            audio_path.unlink(missing_ok=True)
                    except Exception as e:
                        logger.warning(f"Audio extraction/transcription failed for {url}: {e}")
                        
                    # Extract key frames
                    key_frames = []
                    try:
                        key_frames = await self._extract_key_frames(str(local_path))
                    except Exception as e:
                        logger.warning(f"Key frame extraction failed for {url}: {e}")
                        
                    return {
                        'url': url,
                        'local_path': str(local_path),
                        'file_size': video_info.get('size', 0),
                        'content_type': video_info.get('content_type', ''),
                        'duration': metadata.get('duration', 0),
                        'resolution': metadata.get('resolution', ''),
                        'frame_rate': metadata.get('frame_rate', 0),
                        'codec': metadata.get('codec', ''),
                        'audio_transcription': audio_transcription,
                        'key_frames': key_frames,
                        'processed_at': datetime.now().isoformat()
                    }
                    
                except Exception as e:
                    logger.error(f"Failed to process video {url}: {e}")
                    return None
                    
        # Process all videos
        tasks = [process_single_video(url) for url in video_urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out failed results
        for result in results:
            if result and not isinstance(result, Exception):
                processed_videos.append(result)
                
        return processed_videos
        
    async def _process_audio(self, audio_urls: List[str]) -> List[Dict[str, Any]]:
        """Process audio files: download and transcribe."""
        processed_audio = []
        
        semaphore = asyncio.Semaphore(5)
        
        async def process_single_audio(url: str):
            async with semaphore:
                try:
                    # Download audio
                    local_path, audio_info = await self._download_media(url)
                    
                    if not local_path:
                        return None
                        
                    # Transcribe audio
                    transcription_result = await self.transcription_processor.transcribe_audio(str(local_path))
                    
                    # Extract audio metadata
                    metadata = await self._extract_audio_metadata(str(local_path))
                    
                    return {
                        'url': url,
                        'local_path': str(local_path),
                        'file_size': audio_info.get('size', 0),
                        'content_type': audio_info.get('content_type', ''),
                        'duration': metadata.get('duration', 0),
                        'sample_rate': metadata.get('sample_rate', 0),
                        'channels': metadata.get('channels', 0),
                        'codec': metadata.get('codec', ''),
                        'transcription': transcription_result.get('text', ''),
                        'confidence': transcription_result.get('confidence', 0),
                        'language': transcription_result.get('language', ''),
                        'processed_at': datetime.now().isoformat()
                    }
                    
                except Exception as e:
                    logger.error(f"Failed to process audio {url}: {e}")
                    return None
                    
        # Process all audio files
        tasks = [process_single_audio(url) for url in audio_urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out failed results
        for result in results:
            if result and not isinstance(result, Exception):
                processed_audio.append(result)
                
        return processed_audio
        
    async def _process_documents(self, document_urls: List[str]) -> List[Dict[str, Any]]:
        """Process documents: download and extract content."""
        processed_docs = []
        
        semaphore = asyncio.Semaphore(3)
        
        async def process_single_document(url: str):
            async with semaphore:
                try:
                    # Download document
                    local_path, doc_info = await self._download_media(url)
                    
                    if not local_path:
                        return None
                        
                    # Extract content based on file type
                    content = await self._extract_document_content(str(local_path))
                    
                    # Extract any embedded images
                    embedded_images = []
                    if local_path.suffix.lower() == '.pdf':
                        try:
                            embedded_images = await self._extract_images_from_pdf(str(local_path))
                        except Exception as e:
                            logger.warning(f"Failed to extract images from PDF {url}: {e}")
                            
                    return {
                        'url': url,
                        'local_path': str(local_path),
                        'file_size': doc_info.get('size', 0),
                        'content_type': doc_info.get('content_type', ''),
                        'file_type': local_path.suffix.lower(),
                        'content': content.get('text', ''),
                        'metadata': content.get('metadata', {}),
                        'embedded_images': embedded_images,
                        'word_count': len(content.get('text', '').split()),
                        'processed_at': datetime.now().isoformat()
                    }
                    
                except Exception as e:
                    logger.error(f"Failed to process document {url}: {e}")
                    return None
                    
        # Process all documents
        tasks = [process_single_document(url) for url in document_urls]
        results = await asyncio.gather(*tasks, return_exceptions=True)
        
        # Filter out failed results
        for result in results:
            if result and not isinstance(result, Exception):
                processed_docs.append(result)
                
        return processed_docs
        
    async def _download_media(self, url: str) -> Tuple[Optional[Path], Dict[str, Any]]:
        """Download media file and return local path and info."""
        try:
            async with self.session.get(url) as response:
                if response.status != 200:
                    logger.warning(f"Failed to download {url}: HTTP {response.status}")
                    return None, {}
                    
                # Generate filename
                parsed_url = urlparse(url)
                filename = Path(parsed_url.path).name
                if not filename:
                    # Generate filename from URL hash
                    url_hash = hashlib.md5(url.encode()).hexdigest()[:8]
                    content_type = response.headers.get('content-type', '')
                    extension = self._get_extension_from_content_type(content_type)
                    filename = f"media_{url_hash}{extension}"
                    
                # Create unique filename if file exists
                local_path = self.download_dir / filename
                counter = 1
                while local_path.exists():
                    stem = local_path.stem
                    suffix = local_path.suffix
                    local_path = self.download_dir / f"{stem}_{counter}{suffix}"
                    counter += 1
                    
                # Download file
                with open(local_path, 'wb') as f:
                    async for chunk in response.content.iter_chunked(8192):
                        f.write(chunk)
                        
                # Get file info
                file_info = {
                    'size': local_path.stat().st_size,
                    'content_type': response.headers.get('content-type', ''),
                    'last_modified': response.headers.get('last-modified', '')
                }
                
                return local_path, file_info
                
        except Exception as e:
            logger.error(f"Failed to download {url}: {e}")
            return None, {}
            
    def _get_extension_from_content_type(self, content_type: str) -> str:
        """Get file extension from content type."""
        type_map = {
            'image/jpeg': '.jpg',
            'image/png': '.png',
            'image/gif': '.gif',
            'image/webp': '.webp',
            'image/svg+xml': '.svg',
            'video/mp4': '.mp4',
            'video/webm': '.webm',
            'video/quicktime': '.mov',
            'audio/mpeg': '.mp3',
            'audio/wav': '.wav',
            'audio/ogg': '.ogg',
            'application/pdf': '.pdf',
            'application/msword': '.doc',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': '.docx'
        }
        return type_map.get(content_type, '')
        
    async def _analyze_image(self, image_path: Path) -> Dict[str, Any]:
        """Analyze image properties."""
        try:
            from PIL import Image
            import colorsys
            
            with Image.open(image_path) as img:
                analysis = {
                    'dimensions': {
                        'width': img.width,
                        'height': img.height
                    },
                    'format': img.format,
                    'mode': img.mode,
                    'has_transparency': img.mode in ('RGBA', 'LA') or 'transparency' in img.info
                }
                
                # Check for text (simplified heuristic)
                analysis['has_text'] = self._check_image_for_text(img)
                
                # Extract dominant colors
                analysis['dominant_colors'] = self._extract_dominant_colors(img)
                
                return analysis
                
        except Exception as e:
            logger.error(f"Failed to analyze image {image_path}: {e}")
            return {}
            
    def _check_image_for_text(self, img) -> bool:
        """Simple heuristic to check if image might contain text."""
        try:
            # Convert to grayscale for analysis
            if img.mode != 'L':
                img = img.convert('L')
                
            # Simple edge detection to find text-like patterns
            import numpy as np
            from scipy import ndimage
            
            img_array = np.array(img)
            
            # Calculate gradient magnitude
            grad_x = ndimage.sobel(img_array, axis=1)
            grad_y = ndimage.sobel(img_array, axis=0)
            gradient_magnitude = np.sqrt(grad_x**2 + grad_y**2)
            
            # Check for high-frequency content (text-like)
            high_frequency_pixels = np.sum(gradient_magnitude > np.mean(gradient_magnitude) * 2)
            total_pixels = gradient_magnitude.size
            
            return (high_frequency_pixels / total_pixels) > 0.1  # 10% threshold
            
        except ImportError:
            # Fallback without scipy - return True to trigger OCR
            return True
        except Exception:
            # If analysis fails, assume no text
            return False
            
    def _extract_dominant_colors(self, img, num_colors: int = 5) -> List[str]:
        """Extract dominant colors from image."""
        try:
            from PIL import Image
            import numpy as np
            from collections import Counter
            
            # Resize image for faster processing
            img_small = img.copy()
            img_small.thumbnail((150, 150))
            
            # Convert to RGB and get pixel data
            img_array = np.array(img_small.convert('RGB'))
            pixels = img_array.reshape(-1, 3)
            
            # Use k-means-like approach for dominant colors
            # Simple quantization
            quantized = pixels // 32 * 32  # Quantize to 8 levels per channel
            color_counts = Counter([tuple(pixel) for pixel in quantized])
            
            # Get most common colors
            dominant_colors = []
            for color, count in color_counts.most_common(num_colors):
                r, g, b = color
                dominant_colors.append(f"#{r:02x}{g:02x}{b:02x}")
                
            return dominant_colors
            
        except Exception as e:
            logger.error(f"Failed to extract dominant colors: {e}")
            return []
            
    async def _extract_video_metadata(self, video_path: Path) -> Dict[str, Any]:
        """Extract metadata from video file."""
        try:
            import subprocess
            import json
            
            # Use ffprobe to get video metadata
            cmd = [
                'ffprobe', '-v', 'quiet', '-print_format', 'json',
                '-show_format', '-show_streams', str(video_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                metadata = json.loads(result.stdout)
                
                # Find video stream
                video_stream = None
                for stream in metadata.get('streams', []):
                    if stream.get('codec_type') == 'video':
                        video_stream = stream
                        break
                        
                if video_stream:
                    return {
                        'duration': float(metadata.get('format', {}).get('duration', 0)),
                        'resolution': f"{video_stream.get('width', 0)}x{video_stream.get('height', 0)}",
                        'frame_rate': eval(video_stream.get('r_frame_rate', '0/1')),  # Convert fraction
                        'codec': video_stream.get('codec_name', ''),
                        'bitrate': metadata.get('format', {}).get('bit_rate', '')
                    }
                    
        except Exception as e:
            logger.error(f"Failed to extract video metadata: {e}")
            
        return {}
        
    async def _extract_audio_metadata(self, audio_path: Path) -> Dict[str, Any]:
        """Extract metadata from audio file."""
        try:
            import subprocess
            import json
            
            # Use ffprobe to get audio metadata
            cmd = [
                'ffprobe', '-v', 'quiet', '-print_format', 'json',
                '-show_format', '-show_streams', str(audio_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                metadata = json.loads(result.stdout)
                
                # Find audio stream
                audio_stream = None
                for stream in metadata.get('streams', []):
                    if stream.get('codec_type') == 'audio':
                        audio_stream = stream
                        break
                        
                if audio_stream:
                    return {
                        'duration': float(metadata.get('format', {}).get('duration', 0)),
                        'sample_rate': int(audio_stream.get('sample_rate', 0)),
                        'channels': int(audio_stream.get('channels', 0)),
                        'codec': audio_stream.get('codec_name', ''),
                        'bitrate': metadata.get('format', {}).get('bit_rate', '')
                    }
                    
        except Exception as e:
            logger.error(f"Failed to extract audio metadata: {e}")
            
        return {}
        
    async def _extract_audio_from_video(self, video_path: Path) -> Optional[Path]:
        """Extract audio track from video."""
        try:
            import subprocess
            
            audio_path = video_path.with_suffix('.audio.wav')
            
            # Extract audio using ffmpeg
            cmd = [
                'ffmpeg', '-i', str(video_path), '-vn', '-acodec', 'pcm_s16le',
                '-ar', '16000', '-ac', '1', '-y', str(audio_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True)
            if result.returncode == 0:
                return audio_path
            else:
                logger.error(f"Failed to extract audio: {result.stderr}")
                return None
                
        except Exception as e:
            logger.error(f"Audio extraction failed: {e}")
            return None
            
    async def _extract_key_frames(self, video_path: Path, num_frames: int = 5) -> List[Path]:
        """Extract key frames from video."""
        try:
            import subprocess
            
            key_frames = []
            duration = await self._get_video_duration(video_path)
            
            if duration <= 0:
                return key_frames
                
            # Extract frames at regular intervals
            interval = duration / (num_frames + 1)
            
            for i in range(num_frames):
                timestamp = interval * (i + 1)
                frame_path = video_path.with_suffix(f'.frame_{i+1:03d}.jpg')
                
                cmd = [
                    'ffmpeg', '-i', str(video_path), '-ss', str(timestamp),
                    '-vframes', '1', '-q:v', '2', '-y', str(frame_path)
                ]
                
                result = subprocess.run(cmd, capture_output=True)
                if result.returncode == 0:
                    key_frames.append(frame_path)
                    
            return key_frames
            
        except Exception as e:
            logger.error(f"Key frame extraction failed: {e}")
            return []
            
    async def _get_video_duration(self, video_path: Path) -> float:
        """Get video duration in seconds."""
        try:
            import subprocess
            
            cmd = [
                'ffprobe', '-v', 'quiet', '-show_entries', 'format=duration',
                '-of', 'default=noprint_wrappers=1:nokey=1', str(video_path)
            ]
            
            result = subprocess.run(cmd, capture_output=True, text=True)
            if result.returncode == 0:
                return float(result.stdout.strip())
                
        except Exception:
            pass
            
        return 0.0
        
    async def _extract_document_content(self, doc_path: Path) -> Dict[str, Any]:
        """Extract content from document based on file type."""
        try:
            if doc_path.suffix.lower() == '.pdf':
                return await self._extract_pdf_content(str(doc_path))
            elif doc_path.suffix.lower() in ['.docx', '.doc']:
                return await self._extract_word_content(str(doc_path))
            elif doc_path.suffix.lower() == '.txt':
                with open(doc_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                return {
                    'text': content,
                    'metadata': {'encoding': 'utf-8', 'lines': len(content.split('\n'))}
                }
            else:
                # Fallback: try to read as text
                try:
                    with open(doc_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                    return {
                        'text': content,
                        'metadata': {'encoding': 'utf-8', 'type': 'text'}
                    }
                except UnicodeDecodeError:
                    return {'text': '', 'metadata': {'error': 'Unable to decode file'}}
                    
        except Exception as e:
            logger.error(f"Failed to extract document content: {e}")
            return {'text': '', 'metadata': {'error': str(e)}}
            
    async def _extract_pdf_content(self, pdf_path: str) -> Dict[str, Any]:
        """Extract text content from PDF."""
        try:
            import PyPDF2
            
            text = ""
            metadata = {}
            
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                metadata = {
                    'pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else ''
                }
                
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                    
            return {
                'text': text,
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            return {'text': '', 'metadata': {'error': str(e)}}
            
    async def _extract_word_content(self, doc_path: str) -> Dict[str, Any]:
        """Extract content from Word document."""
        try:
            from docx import Document
            
            doc = Document(doc_path)
            
            text = ""
            paragraphs = []
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                paragraphs.append({
                    'text': paragraph.text,
                    'style': paragraph.style.name if paragraph.style else 'Normal'
                })
                
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    table_data.append(row_data)
                tables.append(table_data)
                
            return {
                'text': text,
                'metadata': {
                    'paragraphs': len(paragraphs),
                    'tables': len(tables),
                    'styles': list(set(p.get('style', 'Normal') for p in paragraphs))
                },
                'paragraphs': paragraphs,
                'tables': tables
            }
            
        except Exception as e:
            logger.error(f"Word document extraction failed: {e}")
            return {'text': '', 'metadata': {'error': str(e)}}
            
    async def _extract_images_from_pdf(self, pdf_path: str) -> List[Dict[str, Any]]:
        """Extract images from PDF (basic implementation)."""
        try:
            # This is a simplified implementation
            # For full image extraction from PDF, consider using pdfplumber or similar
            return []
            
        except Exception as e:
            logger.error(f"PDF image extraction failed: {e}")
            return []